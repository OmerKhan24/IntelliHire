"""
HR Document RAG Service - Document Processing & Vector Storage
Handles document upload, text extraction, chunking, and semantic search
Adapted from HRMS advanced_rag_service.py for IntelliHire
"""

import os
import logging
from typing import List, Dict, Any, Optional
from pathlib import Path
import PyPDF2
from docx import Document as DocxDocument
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain.schema import Document
import chromadb
from chromadb.config import Settings
import hashlib
from datetime import datetime

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Process and extract text from various document formats"""
    
    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        """Extract text content from PDF file"""
        try:
            text_content = []
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        text = page.extract_text()
                        if text.strip():
                            text_content.append(text)
                    except Exception as e:
                        logger.warning(f"Error extracting page {page_num} from {file_path}: {e}")
                        continue
            
            full_text = "\n\n".join(text_content)
            logger.info(f"Extracted {len(full_text)} characters from PDF: {file_path}")
            return full_text
            
        except Exception as e:
            logger.error(f"Error reading PDF {file_path}: {e}")
            raise Exception(f"Failed to extract text from PDF: {str(e)}")
    
    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        """Extract text content from DOCX file"""
        try:
            doc = DocxDocument(file_path)
            text_content = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_content.append(row_text)
            
            full_text = "\n\n".join(text_content)
            logger.info(f"Extracted {len(full_text)} characters from DOCX: {file_path}")
            return full_text
            
        except Exception as e:
            logger.error(f"Error reading DOCX {file_path}: {e}")
            raise Exception(f"Failed to extract text from DOCX: {str(e)}")
    
    @staticmethod
    def extract_text_from_txt(file_path: str) -> str:
        """Extract text content from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            logger.info(f"Extracted {len(text)} characters from TXT: {file_path}")
            return text
        except Exception as e:
            logger.error(f"Error reading TXT {file_path}: {e}")
            raise Exception(f"Failed to extract text from TXT: {str(e)}")
    
    @staticmethod
    def extract_text(file_path: str) -> str:
        """Extract text from supported file formats"""
        file_extension = Path(file_path).suffix.lower()
        
        if file_extension == '.pdf':
            return DocumentProcessor.extract_text_from_pdf(file_path)
        elif file_extension in ['.docx', '.doc']:
            return DocumentProcessor.extract_text_from_docx(file_path)
        elif file_extension == '.txt':
            return DocumentProcessor.extract_text_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")


class HRDocumentRAGService:
    """
    RAG Service for HR Documents - handles document processing, 
    chunking, embedding, and semantic search
    """
    
    def __init__(self, upload_folder: str, embedding_model: str = "all-MiniLM-L6-v2"):
        """
        Initialize HR RAG service with ChromaDB
        
        Args:
            upload_folder: Directory where documents are stored
            embedding_model: Sentence transformer model for embeddings
        """
        self.upload_folder = Path(upload_folder)
        self.upload_folder.mkdir(parents=True, exist_ok=True)
        
        # ChromaDB setup (local persistent storage)
        self.chroma_dir = Path(upload_folder).parent / "chroma_db"
        self.chroma_dir.mkdir(parents=True, exist_ok=True)
        
        # Initialize ChromaDB client
        self.chroma_client = chromadb.PersistentClient(
            path=str(self.chroma_dir),
            settings=Settings(
                anonymized_telemetry=False,
                allow_reset=True
            )
        )
        
        self.collection_name = "hr_documents"
        self.embedding_model = embedding_model
        
        # Text splitter for chunking documents
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=500,  # Smaller chunks for precise retrieval
            chunk_overlap=50,
            separators=["\n\n", "\n", ". ", " ", ""],
            length_function=len
        )
        
        # Initialize or get collection
        try:
            self.collection = self.chroma_client.get_or_create_collection(
                name=self.collection_name,
                metadata={"description": "HR policy documents and company information"}
            )
            logger.info(f"✅ ChromaDB collection '{self.collection_name}' initialized")
        except Exception as e:
            logger.error(f"Error initializing ChromaDB collection: {e}")
            raise
    
    def _generate_document_id(self, file_path: str, content: str) -> str:
        """Generate unique ID for document based on path and content"""
        content_hash = hashlib.md5(content.encode()).hexdigest()[:8]
        file_name = Path(file_path).stem
        return f"{file_name}_{content_hash}"
    
    def process_and_store_document(
        self, 
        file_path: str, 
        metadata: Dict[str, Any]
    ) -> Dict[str, Any]:
        """
        Process document: extract text, chunk, embed, and store in vector DB
        
        Args:
            file_path: Path to the uploaded document
            metadata: Document metadata (title, uploaded_by, category, etc.)
        
        Returns:
            Dict with processing results
        """
        try:
            logger.info(f"📄 Processing document: {file_path}")
            
            # Extract text from document
            text_content = DocumentProcessor.extract_text(file_path)
            
            if not text_content or len(text_content.strip()) < 50:
                raise ValueError("Document contains insufficient text content")
            
            # Generate unique document ID
            doc_id = self._generate_document_id(file_path, text_content)
            
            # Check if document already exists
            existing_docs = self.collection.get(
                where={"document_id": doc_id}
            )
            
            if existing_docs and len(existing_docs['ids']) > 0:
                logger.warning(f"⚠️ Document {doc_id} already exists, skipping")
                return {
                    "status": "exists",
                    "document_id": doc_id,
                    "message": "Document already indexed"
                }
            
            # Split text into chunks
            chunks = self.text_splitter.split_text(text_content)
            logger.info(f"📝 Split document into {len(chunks)} chunks")
            
            # Prepare chunks with metadata
            chunk_ids = []
            chunk_texts = []
            chunk_metadatas = []
            
            for i, chunk in enumerate(chunks):
                chunk_id = f"{doc_id}_chunk_{i}"
                chunk_ids.append(chunk_id)
                chunk_texts.append(chunk)
                
                # Add chunk-specific metadata
                chunk_metadata = {
                    **metadata,
                    "document_id": doc_id,
                    "chunk_index": i,
                    "total_chunks": len(chunks),
                    "file_path": file_path,
                    "processed_at": datetime.utcnow().isoformat()
                }
                chunk_metadatas.append(chunk_metadata)
            
            # Add to ChromaDB
            self.collection.add(
                ids=chunk_ids,
                documents=chunk_texts,
                metadatas=chunk_metadatas
            )
            
            logger.info(f"✅ Stored {len(chunks)} chunks for document {doc_id}")
            
            return {
                "status": "success",
                "document_id": doc_id,
                "chunks_created": len(chunks),
                "characters_processed": len(text_content),
                "file_path": file_path
            }
            
        except Exception as e:
            logger.error(f"❌ Error processing document {file_path}: {e}")
            raise Exception(f"Document processing failed: {str(e)}")
    
    def search_documents(
        self, 
        query: str, 
        n_results: int = 5,
        filter_metadata: Optional[Dict[str, Any]] = None
    ) -> List[Dict[str, Any]]:
        """
        Search for relevant document chunks using semantic search
        
        Args:
            query: User query text
            n_results: Number of results to return
            filter_metadata: Optional filters (e.g., category, uploaded_by)
        
        Returns:
            List of relevant document chunks with metadata and scores
        """
        try:
            logger.info(f"🔍 Searching for: '{query}'")
            
            # Build query parameters
            query_params = {
                "query_texts": [query],
                "n_results": n_results
            }
            
            # Add metadata filters if provided
            if filter_metadata:
                query_params["where"] = filter_metadata
            
            # Query ChromaDB
            results = self.collection.query(**query_params)
            
            # Format results
            formatted_results = []
            
            if results and results['documents'] and len(results['documents']) > 0:
                documents = results['documents'][0]
                metadatas = results['metadatas'][0] if results['metadatas'] else []
                distances = results['distances'][0] if results['distances'] else []
                
                for i, doc_text in enumerate(documents):
                    result = {
                        "content": doc_text,
                        "metadata": metadatas[i] if i < len(metadatas) else {},
                        "score": 1 - distances[i] if i < len(distances) else 0  # Convert distance to similarity
                    }
                    formatted_results.append(result)
            
            logger.info(f"✅ Found {len(formatted_results)} relevant chunks")
            return formatted_results
            
        except Exception as e:
            logger.error(f"❌ Search error: {e}")
            raise Exception(f"Search failed: {str(e)}")
    
    def delete_document(self, document_id: str) -> bool:
        """
        Delete document and all its chunks from vector database
        
        Args:
            document_id: Unique document identifier
        
        Returns:
            True if successful
        """
        try:
            # Get all chunks for this document
            results = self.collection.get(
                where={"document_id": document_id}
            )
            
            if not results or not results['ids']:
                logger.warning(f"⚠️ Document {document_id} not found")
                return False
            
            # Delete all chunks
            self.collection.delete(
                ids=results['ids']
            )
            
            logger.info(f"🗑️ Deleted document {document_id} with {len(results['ids'])} chunks")
            return True
            
        except Exception as e:
            logger.error(f"❌ Error deleting document {document_id}: {e}")
            raise Exception(f"Delete failed: {str(e)}")
    
    def get_collection_stats(self) -> Dict[str, Any]:
        """Get statistics about the document collection"""
        try:
            # Get all documents
            all_docs = self.collection.get()
            
            # Count unique documents
            unique_docs = set()
            categories = {}
            
            if all_docs and all_docs['metadatas']:
                for metadata in all_docs['metadatas']:
                    doc_id = metadata.get('document_id')
                    if doc_id:
                        unique_docs.add(doc_id)
                    
                    category = metadata.get('category', 'uncategorized')
                    categories[category] = categories.get(category, 0) + 1
            
            return {
                "total_chunks": len(all_docs['ids']) if all_docs else 0,
                "unique_documents": len(unique_docs),
                "categories": categories,
                "collection_name": self.collection_name
            }
            
        except Exception as e:
            logger.error(f"❌ Error getting collection stats: {e}")
            return {
                "total_chunks": 0,
                "unique_documents": 0,
                "categories": {},
                "error": str(e)
            }
    
    def list_documents(self) -> List[Dict[str, Any]]:
        """List all documents in the collection"""
        try:
            all_docs = self.collection.get()
            
            documents = {}
            
            if all_docs and all_docs['metadatas']:
                for metadata in all_docs['metadatas']:
                    doc_id = metadata.get('document_id')
                    if doc_id and doc_id not in documents:
                        documents[doc_id] = {
                            "document_id": doc_id,
                            "title": metadata.get('title', 'Untitled'),
                            "category": metadata.get('category', 'uncategorized'),
                            "uploaded_by": metadata.get('uploaded_by', 'unknown'),
                            "processed_at": metadata.get('processed_at'),
                            "file_path": metadata.get('file_path')
                        }
            
            return list(documents.values())
            
        except Exception as e:
            logger.error(f"❌ Error listing documents: {e}")
            return []
