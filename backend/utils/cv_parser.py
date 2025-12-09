"""
CV Parser Utility
Extracts text from PDF and DOC/DOCX files for RAG processing
"""
import os
import logging
from typing import Optional

logger = logging.getLogger(__name__)

def extract_text_from_cv(file_path: str) -> Optional[str]:
    """
    Extract text content from CV file (PDF, DOC, DOCX)
    
    Args:
        file_path: Path to the CV file
        
    Returns:
        Extracted text content or None if extraction fails
    """
    if not os.path.exists(file_path):
        logger.error(f"CV file not found: {file_path}")
        return None
    
    file_ext = file_path.lower().split('.')[-1]
    
    try:
        if file_ext == 'pdf':
            return _extract_from_pdf(file_path)
        elif file_ext in ['doc', 'docx']:
            return _extract_from_docx(file_path)
        else:
            logger.warning(f"Unsupported file type: {file_ext}")
            return None
    except Exception as e:
        logger.error(f"Failed to extract text from {file_path}: {e}")
        return None


def _extract_from_pdf(file_path: str) -> str:
    """Extract text from PDF file using PyPDF2"""
    try:
        from PyPDF2 import PdfReader
        
        reader = PdfReader(file_path)
        text_content = []
        
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_content.append(text)
        
        extracted_text = '\n'.join(text_content)
        logger.info(f"✅ Extracted {len(extracted_text)} characters from PDF")
        return extracted_text
        
    except ImportError:
        logger.error("PyPDF2 not installed. Run: pip install PyPDF2")
        return None
    except Exception as e:
        logger.error(f"PDF extraction failed: {e}")
        return None


def _extract_from_docx(file_path: str) -> str:
    """Extract text from DOCX file using python-docx"""
    try:
        import docx
        from docx import Document
        
        doc = Document(file_path)
        text_content = []
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_content.append(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_content.append(cell.text)
        
        extracted_text = '\n'.join(text_content)
        logger.info(f"✅ Extracted {len(extracted_text)} characters from DOCX")
        return extracted_text
        
    except ImportError as ie:
        logger.error(f"python-docx import error: {ie}. Run: pip install python-docx")
        return None
        return None
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        return None


def summarize_cv_for_context(cv_text: str, max_length: int = 2000) -> str:
    """
    Create a concise summary of CV for AI context
    Extracts key information like skills, experience, education
    
    Args:
        cv_text: Full CV text
        max_length: Maximum characters to return
        
    Returns:
        Summarized CV text suitable for AI prompts
    """
    if not cv_text:
        return ""
    
    # If CV is already short enough, return as is
    if len(cv_text) <= max_length:
        return cv_text
    
    # Simple truncation with ellipsis (could be enhanced with NLP summarization)
    return cv_text[:max_length] + "..."
